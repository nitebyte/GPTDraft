import openai
import os
import csv
import requests
import json
import sys
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def define_API(key):
    openai.api_key = key
    

def PR(prompt, system_content="", token=24, temperature=0.35, pp=0.15, fp=0.05):
    try:
        # Create a list of messages containing the system content and user prompt
        if len(system_content) > 0:
            messages = [{'role': 'system', 'content': system_content}, {'role': 'user', 'content': prompt}]
        else:
            messages = [{'role': 'user', 'content': prompt}]
        

        # Call the OpenAI API to generate a chat completion given the messages and parameters
        response = openai.ChatCompletion.create(model="gpt-3.5-turbo", messages=messages, temperature=temperature, max_tokens=token, presence_penalty=pp, frequency_penalty=fp)

        # Get the system message from the API response
        system_message = response.choices[0].get('message', {}).get('content', '').strip()

        # Get the number of tokens used from the API response
        tokens = response.get('usage', {}).get('total_tokens', 0)

        # Calculates the cost of the API response
        cost = tokens * 0.000002

        # Get the user prompt message and remove any leading or trailing whitespace
        prompt_message = prompt.strip()

        # Return a dictionary containing the status of the message, the number of tokens used, and the system message
        return {"status": 1, "tokens": tokens, "cost": cost, "content": system_message}
    except Exception as e:
        # If an exception occurs, return a dictionary containing the status of the message and the error message
        return {"status": 0, "tokens":0, "cost": "", "content": str(e)}



def append_text_to_file(text, file_name):
    # Extract the folder name by removing the file extension
    folder_name = os.path.splitext(file_name)[0]
    
    # Create the folder if it does not exist
    if not os.path.exists(folder_name):
        os.makedirs(folder_name)

    # Open a file in "append+" mode, which creates the file if it does not exist and appends new content to the end of the file
    with open(os.path.join(folder_name, file_name), "a+", encoding="utf-8") as f:
        # Write the given text to the end of the file, followed by a newline character
        text += '\n'
        f.write(text)



def bookArray(chapterListRaw):
    chapterList = chapterListRaw.split('\n')

    book = [[[None]]]

    for chapter in chapterList:
        if not chapter.strip():  # Check if the line is empty or contains only whitespace
            continue              # Skip empty lines
        parts = chapter.split(' ')
        chapterNum = int(parts[0].split('.')[0])
        sectionNum = int(parts[0].split('.')[1])
        partNum = int(parts[0].split('.')[2])
        title = ' '.join(parts[1:])
  
        while len(book) < chapterNum+1:
            book.append([[[None]]])
        while len(book[chapterNum]) < sectionNum+1:
            book[chapterNum].append([[[None]]])
        while len(book[chapterNum][sectionNum]) < partNum+1:
            book[chapterNum][sectionNum].append([None])
  
        book[chapterNum][sectionNum][partNum] = title
    return book



def print_book(book, pre=""):
    for chapterNum, chapter in enumerate(book):
        for sectionNum, section in enumerate(chapter):
            for partNum, part in enumerate(section):
                if part is not None:  # Check if part is not empty
                    space = ""
                    if sectionNum > 0:
                        space += "  "
                    if partNum > 0:
                        space += "  "
                    print(f"{pre}{space}{chapterNum}.{sectionNum}.{partNum} {part}")





def serialize_book(book):
    serialized = ""
    for chapterNum, chapter in enumerate(book):
        for sectionNum, section in enumerate(chapter):
            for partNum, part in enumerate(section):
                if part:
                    serialized += f"{part},"
    return serialized.rstrip(",")  # Remove the trailing comma before returning



def count_parts(book):
    num = 0
    for chapterNum, chapter in enumerate(book):
        for sectionNum, section in enumerate(chapter):
            for partNum, part in enumerate(section):
                if part:
                    num+=1
    return num


def Title(text, place):
    return text[6:]

def txt_to_docx(txt_file, docx_file,bookSubject,api,d_sec,serialBook,documentType):
    # Prompt the user to enter the subject and document type
    print("╠═════════════════════════════════════════════════════════════════════════════════")
    ibookName = input("║ Enter The Name Of The Book: ")
    print("╠═════════════════════════════════════════════════════════════════════════════════")
    bookSubtitle = input("║ Enter The Subtitle (Or Leave Blank To Auto-Generate): ")
    print("╠═════════════════════════════════════════════════════════════════════════════════")
    ibookAuthor = input("║ Enter The Author Name: ")
    print("╠═════════════════════════════════════════════════════════════════════════════════")
    ibookHeight = 9
    ibookWidth = 6
    ibookWidth = input("║ Enter The Book Width In Inches: ")
    print("╠═════════════════════════════════════════════════════════════════════════════════")
    ibookHeight = input("║ Enter The Book Height In Inches: ")

    if len(bookSubtitle)<1:
        bookSubtitle = PR("Generate a short subtitle for a book titled " + bookSubject + " and is the following type of book:" + documentType + ". The chapters in the book are: " + serialBook, "", 100, 0.75)['content']
    #Set URL for DALL-E
    url = "https://api.openai.com/v1/images/generations"
    print("║ Starting Conversion")
    # Create a new Document
    document = Document()
    print("║ .docx Created")
    # Set the page size and margins
    section = document.sections[0]
    section.page_width = Inches(ibookWidth)
    section.page_height = Inches(ibookHeight)
    section.left_margin = Inches(0.375)
    section.right_margin = Inches(0.375)
    section.top_margin = Inches(0.25)
    section.bottom_margin = Inches(0.25)
    print("║ Page Size & Margins Set")


    # Add half-title page
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(bookSubject)
    title_run.font.size = Pt(18)
    title_run.bold = True
    document.add_page_break()
    print("║ Title Page Created")


     # Add title page
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(bookSubject)
    title_run.font.size = Pt(20)
    title_run.bold = True
    title_paragraph.add_run("\n\n" + bookSubtitle).italic = True
    title_paragraph.add_run("\n\n" + ibookAuthor + "").bold = True
    document.add_page_break()
    print("║ Title Page Created")

    # Add copyright page
    copyright_paragraph = document.add_paragraph()
    copyright_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    copyright_paragraph.add_run("Copyright © 2023 by " + ibookAuthor + "\n\nAll rights reserved.\n\nThis publication aims to provide factual and authoritative information on the topic discussed. It is important to note that neither the author nor the publisher provides legal, investment, accounting, or any other professional services. Although the author and publisher have taken every effort to ensure the accuracy and comprehensiveness of the content presented, they do not make any express or implied warranties of merchantability or fitness for a particular purpose. The information provided in this book may not be suitable for every situation, and readers are advised to seek professional advice as necessary. The author and publisher will not be held liable for any financial or other damages resulting from the use of this information.\n\nCover art and illustrations by DALL-E 2\n\nProofreading provided by GPT4\n\nEditing services by " + ibookAuthor + "\n\nGPT/DALL-E 2 integration software developed by " + ibookAuthor + "").font.size = Pt(8)
    print("║ Copyright Page Added")
    document.add_page_break()

    # Add dedication page
    dedication_paragraph = document.add_paragraph()
    dedication_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dedication_paragraph.add_run("Dedication\n\n")
    dedication_paragraph.add_run(PR("Craft a generic one-paragraph dedication for a book about " + bookSubject + " that is authored by a single individual. Please consider using language that would apply universally to readers of the book, rather than addressing any particular individual or group of people.", "", 1024, 0.5)['content']).italic = True
    document.add_page_break()
    print("║ Dedication Page Added")

    # Set the folder name
    folder_name = os.path.splitext(txt_file)[0]

    # Read the input text file from inside the folder
    with open(os.path.join(folder_name, txt_file), 'r', encoding='utf-8') as f:
        txt_lines = f.readlines()
        txt_lines = [line for line in txt_lines if line.strip()]    

    # Process the text file lines
    for line in txt_lines:
        if line.startswith('h1'):
            print("║ Detected Heading Level 1")
            document.add_page_break()
            title = line[2:].strip()
            heading = document.add_heading(level=1)
            heading_run = heading.add_run(title + "\n")
            heading_run.font.size = Pt(18)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer " + api
            }
            imgprompt = PR("I want you to write a DALL-E image generation prompt to generate a photo that represents a chapter of a book about " + title + " in regards to " + bookSubject + ". For example, if you were asked to write a prompt for an image about the geography of Damascus, Virginia, you may output: Mountainous terrain with dense forests and trees, peaceful and serene, in the vicinity of Damascus, VA, USA. Shot on a Canon EOS R6 with a Canon RF 24-105mm f/4L IS USM Lens, 4K film still, natural lighting, vibrant colors, crisp details, and soft shadows. There is a hard limit of 300 characters in your responce - any longer and it will be cut off.", "", 100, 0.5)
            print("║ Generating Image For: " + title + " : " + bookSubject)
            data = {
                "prompt": imgprompt["content"],
                "n": 1,
                "size": "1024x1024"
            }
            #print(imgprompt)
            response = requests.post(url, headers=headers, data=json.dumps(data))

            if response.status_code == 200:
                print("Good Image")
                result = response.json()
                # Download the image from a link
                image_url = result['data'][0]['url']
                response = requests.get(image_url)
                img_content = response.content

                # Create the folder if it doesn't exist
                folder_name = bookSubject
                try:
                    if not os.path.exists(folder_name):
                        os.makedirs(folder_name)

                    # Save the image to the folder
                    with open(os.path.join(folder_name, f"{title} - {bookSubject}.png"), "wb") as f:
                        f.write(img_content)

                    img = BytesIO(img_content)

                except Exception as e:
                    # Handle the exception
                    print(f"Error saving image: {e}")
                # Add a paragraph with an image
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # center the paragraph
                run = paragraph.add_run()
                run.add_picture(img, width=Inches(4))  # adjust width as necessary
                #document.add_page_break()
            else:
                print("Bad Image")

        elif line.startswith('h2'):
            print("║ Detected Heading Level 2")
            title = line[2:].strip()
            heading = document.add_heading(level=2)
            heading_run = heading.add_run(title + "\n")
            heading_run.font.size = Pt(16)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if d_sec == "Y":
                headers = {
                    "Content-Type": "application/json",
                    "Authorization": f"Bearer " + api
                }
                imgprompt = PR("I want you to write a DALL-E image generation prompt to generate a photo that represents a chapter of a book about " + title + " in regards to " + bookSubject + ". For example, if you were asked to write a prompt for an image about the geography of Damascus, Virginia, you may output: Mountainous terrain with dense forests and trees, peaceful and serene, in the vicinity of Damascus, VA, USA. Shot on a Canon EOS R6 with a Canon RF 24-105mm f/4L IS USM Lens, 4K film still, natural lighting, vibrant colors, crisp details, and soft shadows. There is a hard limit of 300 characters in your responce - any longer and it will be cut off.", "", 100, 0.5)
                print("║ Generating Image For: " + title + " : " + bookSubject)
                data = {
                    "prompt": imgprompt["content"],
                    "n": 1,
                    "size": "1024x1024"
                }
                #print(imgprompt)

                response = requests.post(url, headers=headers, data=json.dumps(data))

                if response.status_code == 200:
                    print("Good Image")
                    result = response.json()
                    # Download the image from a link
                    image_url = result['data'][0]['url']
                    response = requests.get(image_url)
                    img_content = response.content

                    # Create the folder if it doesn't exist
                    folder_name = bookSubject
                    try:
                        if not os.path.exists(folder_name):
                            os.makedirs(folder_name)

                        # Save the image to the folder
                        with open(os.path.join(folder_name, f"{title} - {bookSubject}.png"), "wb") as f:
                            f.write(img_content)

                        img = BytesIO(img_content)

                    except Exception as e:
                        # Handle the exception
                        print(f"Error saving image: {e}")
                    # Add a paragraph with an image
                    paragraph = document.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # center the paragraph
                    run = paragraph.add_run()
                    run.add_picture(img, width=Inches(4))  # adjust width as necessary
                    #document.add_page_break()
                else:
                    print("Bad Image")

        elif line.startswith('h3'):
            print("║ Detected Heading Level 3 - NoImg")
            title = line[2:].strip()
            heading = document.add_heading(level=3)
            heading_run = heading.add_run(title)
            heading_run.font.size = Pt(14)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        else:
            content = line.strip()
            paragraph = document.add_paragraph(content)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Save the DOCX file to the folder
    docx_file = os.path.join(folder_name, bookSubject + ".docx")
    document.save(docx_file)



